"""Multi-format document parser — PDF, DOCX, MD, TXT."""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Output model
# ---------------------------------------------------------------------------

@dataclass
class ParsedDocument:
    """Unified parse result regardless of input format."""

    title: str
    content: str
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    file_path: str = ""

    def __repr__(self) -> str:
        preview = self.content[:80].replace("\n", " ")
        return (
            f"ParsedDocument(title={self.title!r}, "
            f"chars={len(self.content)}, pages={self.page_count}, "
            f"preview={preview!r}...)"
        )


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF (fitz) is required for PDF parsing. pip install PyMuPDF")

    doc = fitz.open(str(path))
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text("text") or "")
    doc.close()

    content = "\n\n".join(pages)
    return ParsedDocument(
        title=path.stem,
        content=content.strip(),
        metadata={"format": "pdf", "original_path": str(path)},
        page_count=len(pages),
        file_path=str(path),
    )


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. pip install python-docx")

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    return ParsedDocument(
        title=path.stem,
        content=content.strip(),
        metadata={"format": "docx", "original_path": str(path)},
        page_count=0,  # DOCX has no fixed page count
        file_path=str(path),
    )


def _parse_text(path: Path) -> ParsedDocument:
    content = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(
        title=path.stem,
        content=content.strip(),
        metadata={"format": path.suffix.lstrip("."), "original_path": str(path)},
        file_path=str(path),
    )


# ---------------------------------------------------------------------------
# Unified API
# ---------------------------------------------------------------------------

_PARSER_MAP = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".doc": _parse_docx,
    ".md": _parse_text,
    ".txt": _parse_text,
    ".markdown": _parse_text,
    ".rst": _parse_text,
    ".py": _parse_text,
    ".json": _parse_text,
    ".yaml": _parse_text,
    ".yml": _parse_text,
    ".csv": _parse_text,
}


def parse_file(file_path: str | Path) -> ParsedDocument | None:
    """Parse a single file into a ParsedDocument.

    Returns None if the format is unsupported or the file is unreadable.
    """
    path = Path(file_path)
    if not path.exists():
        logger.error("File not found: %s", path)
        return None

    suffix = path.suffix.lower()
    parser = _PARSER_MAP.get(suffix)
    if parser is None:
        logger.warning("Unsupported file format: %s (%s)", path, suffix)
        return None

    try:
        doc = parser(path)
        logger.info("Parsed %s: %d chars, %d pages", path.name, len(doc.content), doc.page_count)
        return doc
    except Exception:
        logger.exception("Failed to parse %s", path)
        return None
